from docx import Document
document=Document()
document.save('test.docx')         #新建一个空白文档
document.add_paragraph('添加段落')       #添加段落
document.add_heading('标题')                      #添加标题
document.add_page_break()          #添加分页
table=document.add_table(3,3)      #添加表格，三行三列的表格
row=table.rows[1]                  #第二行
row.cells[0].text='python'
row.cells[1].text='JAVA'
row.cells[2].text='JavaScript'
document.save('test.docx')